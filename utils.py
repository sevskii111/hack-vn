import re
from pathlib import Path
from multiprocessing import Pool
from tqdm import tqdm
from docx.enum.text import WD_COLOR_INDEX
import docx

FACTORS = ['3а', '3б', '3в', '3г', '3д',
           '3е', '3ж', '3з', '3и', '4а', '4б', '4в']


def parse_npa(doc):
    docs = []
    curr_section = None
    curr_doc = None

    for p in doc.paragraphs:
        if p.text == "ОБЩИЕ ПОЛОЖЕНИЯ":
            p.text = "1. ОБЩИЕ ПОЛОЖЕНИЯ"
        section = re.search(r'^((?:\d+\.)+[\s\tА-Я])', p.text.strip())
        if curr_section is not None and section is None:
            if len(p.text.strip()) > 0:
                curr_doc[curr_section].append(p.text.strip())
            continue
        if curr_section is None and section is None:
            continue
        section_text = re.match(
            r'^(?:\d+\.)+', section.group(0).strip()).group(0)
        if re.match(r'^1.\s?$', section_text):
            if curr_doc is not None:
                docs.append(curr_doc)
            curr_doc = {}

        curr_section = section_text
        if curr_doc is None:
            curr_doc = {}
        curr_doc[curr_section] = [p.text.strip()]
    docs.append(curr_doc)

    return docs


def get_dataset():
    docs_train = list(Path().glob('DataSet_*/**/Expertise_Text.docx'))
    return docs_train


def find_corruption_extended_new(doc, possible_points):
    result = []
    possible_points_striped = [pp.strip(' .') for pp in possible_points]
    doc_point = None
    for p in doc.paragraphs:
        doc_points_m = re.findall(
            r'\s\d+\.?(?:\d+[\.\s]?)*\s(?!Методики)(?!«)(?!")', p.text)
        for doc_point_m in doc_points_m:
            doc_point_m = doc_point_m.strip(' .')
            if doc_point_m in possible_points_striped:
                doc_point = doc_point_m

        if re.findall(r'[^\d][34]\sМетодики', p.text):
            point = re.search(r'[34]', p.text).group(0)
            if point == '3':
                sub_points = re.findall(r'\s[«"]?[абвджи][»"\)]', p.text)
            else:
                sub_points = re.findall(r'\s[«"]?[а-в][»"\)]', p.text)
            if len(sub_points) > 0:
                sub_points = [re.search(r'[а-и]', p).group(0)
                              for p in sub_points]
                result.append((point, sub_points, doc_point))

    return result


def parse_doc(expertise_path):
    try:
        edition_path = expertise_path.parent / 'Edition_Text.docx'
        if not edition_path.is_file():
            return None
        edition = docx.Document(edition_path)
        parsed_npa = parse_npa(edition)
        if None not in parsed_npa:
            # print(parsed_npa)
            possible_points = [npa.keys() for npa in parsed_npa]
            possible_points = [
                item for sublist in possible_points for item in sublist]
        else:
            return None
        # print(possible_points)
        expertise = docx.Document(expertise_path)
        # print(expertise_path)
        res = find_corruption_extended_new(expertise, possible_points)
        if len(res) < 1:
            return None
        # for a in res:
        #     total += 1
        # if a[2] is None:
        #     bad += 1
        return (expertise_path, edition_path, res)

    except Exception as e:
        print(e)
        return None
    # print(res)


def parse_docs(docs):
    with Pool(32) as p:
        return list(tqdm(p.imap(parse_doc, docs), total=len(docs)))


def generate_gt(r):
    result = []
    result_keys = ['3а', '3б', '3в', '3г', '3д',
                   '3е', '3ж', '3з', '3и', '4а', '4б', '4в']
    for e in r:
        if e is None:
            continue
        doc_result = [0] * len(result_keys)
        for f in e[2]:
            for pp in f[1]:
                s = f[0] + pp
                doc_result[result_keys.index(s)] = 1
        result.append(doc_result)
    return result


def npa_highlight(doc_path, factors=None, path=None):
    if factors is None:
        expertise = parse_doc(doc_path)
        npa = docx.Document(expertise[1])
        parsed_expertise = parse_doc(expertise[0])
        factors = parsed_expertise[2]
    else:
        npa = docx.Document(doc_path)
    for p in npa.paragraphs:
        for f in factors:
            if re.match(f'^{re.escape(f[2])}[\s\tА-Я]', p.text):
                p.text += f'({f[0]}_{",".join(f[1])})'
                for run in p.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if path is None:
        path = 'npa_highlighted.docx'
    npa.save(path)
